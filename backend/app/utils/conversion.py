import os
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from typing import Optional
from ..core.config import settings
from ..utils.logger import logger


class Conversion:
    """
    Utility class for converting resume files between different formats.
    Supports conversions from ODT, DOCX, and PDF to Markdown and HTML.
    """

    # Class variable for resume directory path
    RESUME_DIR = settings.resume_dir

    @classmethod
    def _get_file_path(cls, file_name: str) -> Path:
        """
        Helper method to construct full file path.

        Args:
            file_name: Name of the file

        Returns:
            Path object for the file
        """
        return Path(cls.RESUME_DIR) / file_name

    @classmethod
    def _get_convertapi_key(cls) -> Optional[str]:
        """
        Retrieve ConvertAPI key from database.

        Returns:
            API key string if found and valid, None otherwise
        """
        from ..core.database import SessionLocal
        from sqlalchemy import text

        db = SessionLocal()
        try:
            query = text("SELECT convertapi_key FROM personal LIMIT 1")
            result = db.execute(query).first()
            if not result or not result.convertapi_key:
                logger.error("ConvertAPI key not found in database")
                return None
            api_key = result.convertapi_key.strip() if result.convertapi_key else None
            if not api_key:
                logger.error("ConvertAPI key is empty after stripping")
                return None
            return api_key
        finally:
            db.close()

    @classmethod
    def _set_file(cls, company: str, title: str, mimetype: str) -> str:
        """
        This will create a filename using the values from the company name and job title

        :param company: Company name
        :param title: Job position title
        :param mimetype: File format as the file extension to use
        :return: full path and filename with extension
        """
        filetmp = company.strip() + '-' + title.strip()
        filename = filetmp.replace(' ', '_') + '.' + mimetype
        return filename


    @classmethod
    def rename_file(cls, filename: str, mime_type: str) -> str:
        """
        Rename a file by replacing its extension with the specified mime_type.

        Args:
            filename: Original filename (e.g., "resume.pdf")
            mime_type: Target mime type/extension (e.g., "docx", "odt", "pdf")

        Returns:
            New filename with updated extension (e.g., "resume.docx")
        """
        # Remove existing extension
        if '.' in filename:
            base_name = filename.rsplit('.', 1)[0]
        else:
            base_name = filename

        # Add new extension
        return f"{base_name}.{mime_type}"

    @classmethod
    def pageFormatting(cls, filename: str, name: str) -> bool:
        """
        Do final formatting on docx file
        - changes page margins to 1/2"
        - center justifies text using Heading 1, Heading 2 or Heading 3 (with name being excluded)
        - changes the font to use Liberation Sans

        Returns:
            boolean - true for success and false for failure
        """
        target_styles = ['Heading 1', 'Heading 2', 'Heading 3']
        margin = 1.3
        document = Document(str(cls._get_file_path(filename)))

        try:
            # change document default font
            style = document.styles['Normal']
            font = style.font
            font.name = 'Liberation Sans'

            # changing the page margins
            sections = document.sections
            for section in sections:
                section.top_margin = Cm(margin)
                section.bottom_margin = Cm(margin)
                section.left_margin = Cm(margin)
                section.right_margin = Cm(margin)

            for paragraph in document.paragraphs:
                if paragraph.text == name:
                    print(f"Matched name: {paragraph.text}")

                elif paragraph.style.name in target_styles:
                    # If it is, set the alignment to center
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"Centered paragraph with style: {paragraph.style.name}")

            document.save(str(cls._get_file_path(filename)))
            return True

        except Exception as e:
            print(f"Error while editing file: {e}")
            return False


    @classmethod
    def _markdown_to_html(cls, markdown_content: str) -> str:
        """
        Convert Markdown content to HTML.

        Args:
            markdown_content: Markdown text

        Returns:
            HTML content as string
        """
        try:
            import markdown

            # Convert markdown to HTML with common extensions
            html_content = markdown.markdown(
                markdown_content,
                extensions=[
                    'extra',          # Tables, fenced code blocks, etc.
                    'nl2br',          # Newline to <br>
                    'sane_lists',     # Better list handling
                    'codehilite',     # Syntax highlighting for code
                    'toc'             # Table of contents
                ]
            )

            # Wrap in basic HTML structure
            html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif;
            line-height: 1.6;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 24px;
            margin-bottom: 16px;
            font-weight: 600;
            line-height: 1.25;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 16px 0;
        }}
        table th, table td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        table th {{
            background-color: #f2f2f2;
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 12px;
            border-radius: 6px;
            overflow-x: auto;
        }}
        pre code {{
            background-color: transparent;
            padding: 0;
        }}
        blockquote {{
            border-left: 4px solid #ddd;
            margin: 0;
            padding-left: 16px;
            color: #666;
        }}
        ul, ol {{
            padding-left: 24px;
        }}
        a {{
            color: #0066cc;
            text-decoration: none;
        }}
        a:hover {{
            text-decoration: underline;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

            return html

        except Exception as e:
            raise Exception(f"Error converting Markdown to HTML: {str(e)}")

    # ODT conversions
    @classmethod
    def odtToMd(cls, file_name: str) -> str:
        """
        Convert ODT file to Markdown.

        Args:
            file_name: Name of the ODT file

        Returns:
            Markdown content as string
        """
        try:
            import subprocess

            file_path = cls._get_file_path(file_name)

            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_name}")

            # Use pandoc for ODT to Markdown conversion
            result = subprocess.run(
                ['pandoc', str(file_path), '-f', 'odt', '-t', 'markdown'],
                capture_output=True,
                text=True,
                check=True
            )

            return result.stdout
        except subprocess.CalledProcessError as e:
            raise Exception(f"Error converting ODT to Markdown: {e.stderr}")
        except Exception as e:
            raise Exception(f"Error converting ODT to Markdown: {str(e)}")

    @classmethod
    def odtToHtml(cls, file_name: str) -> str:
        """
        Convert ODT file to HTML.

        Args:
            file_name: Name of the ODT file

        Returns:
            HTML content as string
        """
        try:
            import subprocess

            file_path = cls._get_file_path(file_name)

            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_name}")

            # Use pandoc for ODT to HTML conversion
            result = subprocess.run(
                ['pandoc', str(file_path), '-f', 'odt', '-t', 'html'],
                capture_output=True,
                text=True,
                check=True
            )

            return result.stdout
        except subprocess.CalledProcessError as e:
            raise Exception(f"Error converting ODT to HTML: {e.stderr}")
        except Exception as e:
            raise Exception(f"Error converting ODT to HTML: {str(e)}")

    # DOCX conversions
    @classmethod
    def docxToMd(cls, file_name: str) -> str:
        """
        Convert DOCX file to Markdown.

        Args:
            file_name: Name of the DOCX file

        Returns:
            Markdown content as string
        """
        try:
            from docx2md.convert import do_convert

            file_path = cls._get_file_path(file_name)

            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_name}")

            # Convert DOCX to Markdown
            markdown_content = do_convert(str(file_path))

            return markdown_content
        except Exception as e:
            raise Exception(f"Error converting DOCX to Markdown: {str(e)}")

    @classmethod
    def docxToHtml(cls, file_name: str) -> str:
        """
        Convert DOCX file to HTML.

        Args:
            file_name: Name of the DOCX file

        Returns:
            HTML content as string
        """
        try:
            from docx_parser_converter.docx_to_html.docx_to_html_converter import DocxToHtmlConverter
            from docx_parser_converter.docx_parsers.utils import read_binary_from_file_path

            file_path = cls._get_file_path(file_name)

            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_name}")

            docx_file_content = read_binary_from_file_path(file_path)

            converter = DocxToHtmlConverter(docx_file_content, use_default_values=True)
            html_output = converter.convert_to_html()

            return html_output

        except Exception as e:
            raise Exception(f"Error converting DOCX to HTML: {str(e)}")

    # PDF conversions using MarkItDown
    @classmethod
    def pdfToMd(cls, file_name: str) -> str:
        """
        Convert PDF file to Markdown using MarkItDown.
        Preserves document structure including headings, lists, tables, and formatting.

        Args:
            file_name: Name of the PDF file

        Returns:
            Markdown content as string
        """
        try:
            from markitdown import MarkItDown

            file_path = cls._get_file_path(file_name)

            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_name}")

            # Initialize MarkItDown converter
            md = MarkItDown()

            # Convert to markdown
            result = md.convert(str(file_path))

            return result.text_content

        except FileNotFoundError:
            raise
        except Exception as e:
            raise Exception(f"Error converting PDF to Markdown: {str(e)}")

    @classmethod
    def pdfToHtml(cls, file_name: str) -> str:
        """
        Convert PDF file to HTML.
        First converts to Markdown using MarkItDown, then converts Markdown to HTML.

        Args:
            file_name: Name of the PDF file

        Returns:
            HTML content as string
        """
        try:
            # Convert PDF to Markdown first
            markdown_content = cls.pdfToMd(file_name)

            # Convert Markdown to HTML
            return cls._markdown_to_html(markdown_content)

        except Exception as e:
            raise Exception(f"Error converting PDF to HTML: {str(e)}")

    # Markdown conversions
    @classmethod
    def mdToHtml(cls, markdown_content: str) -> str:
        """
        Convert Markdown content to HTML.

        Args:
            markdown_content: Markdown text content

        Returns:
            HTML content as string
        """
        try:
            return cls._markdown_to_html(markdown_content)
        except Exception as e:
            raise Exception(f"Error converting Markdown to HTML: {str(e)}")

    @classmethod
    def md2docx(cls, markdown_content: str, output_file_name: str, reference_file: Optional[str] = None) -> str:
        """
        Convert Markdown content to DOCX format.

        Args:
            markdown_content: Markdown text content
            output_file_name: Name for the output DOCX file
            reference_file: Optional reference DOCX file for styling (filename only, not full path)

        Returns:
            Path to the created DOCX file
        """
        try:
            import subprocess
            import tempfile

            # Create temporary markdown file
            with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as temp_md:
                temp_md.write(markdown_content)
                temp_md_path = temp_md.name

            # Output path for docx
            output_path = cls._get_file_path(output_file_name)
            # Ensure directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            try:
                # Build pandoc command
                pandoc_cmd = ['pandoc', temp_md_path, '-f', 'markdown', '-t', 'docx', '-o', str(output_path)]

                # Add reference doc if provided
                if reference_file:
                    logger.debug(f"Using reference file: {reference_file}")

                    reference_path = cls._get_file_path(reference_file)
                    if os.path.exists(reference_path):
                        pandoc_cmd.extend(['--reference-doc', str(reference_path)])

                # Convert markdown to DOCX using pandoc
                result = subprocess.run(
                    pandoc_cmd,
                    capture_output=True,
                    text=True,
                    check=True
                )

                # Clean up temporary file
                os.unlink(temp_md_path)

                return str(output_path)

            except subprocess.CalledProcessError as e:
                # Clean up temporary file on error
                if os.path.exists(temp_md_path):
                    os.unlink(temp_md_path)
                raise Exception(f"Pandoc conversion to DOCX failed: {e.stderr}")

        except Exception as e:
            raise Exception(f"Error converting Markdown to DOCX: {str(e)}")

    # Markdown to DOCX conversion
    @classmethod
    def md2docx_from_job(cls, job_id: int, db) -> dict:
        """
        Convert resume markdown to DOCX format for a specific job.

        This method:
        1. Retrieves the resume_md_rewrite for the job's associated resume
        2. Converts the markdown to DOCX format using pandoc
        3. Saves the DOCX file with a -tailored.docx suffix
        4. Updates the resume_detail.rewrite_file_name field
        5. Returns the new filename

        Args:
            job_id: ID of the job
            db: Database session

        Returns:
            Dictionary with file_name key
        """
        try:
            import subprocess
            import tempfile
            from sqlalchemy import text

            # Query to get resume data
            query = text("""
                SELECT rd.resume_md_rewrite, r.file_name, r.resume_id
                FROM job j
                JOIN resume r ON (j.resume_id = r.resume_id)
                JOIN resume_detail rd ON (r.resume_id = rd.resume_id)
                WHERE j.job_id = :job_id
            """)

            result = db.execute(query, {"job_id": job_id}).first()

            if not result:
                raise ValueError(f"No resume found for job_id {job_id}")

            if not result.resume_md_rewrite:
                raise ValueError(f"No rewritten markdown content found for job_id {job_id}")

            # Extract values
            md_content = result.resume_md_rewrite
            original_file_name = result.file_name
            resume_id = result.resume_id

            # Remove file extension from original filename
            if '.' in original_file_name:
                base_name = original_file_name.rsplit('.', 1)[0]
            else:
                base_name = original_file_name

            # Create new filename with -tailored.docx suffix
            new_file_name = f"{base_name}-tailored.docx"

            # Create temporary markdown file
            with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as temp_md:
                temp_md.write(md_content)
                temp_md_path = temp_md.name

            # Output path for docx
            output_path = cls._get_file_path(new_file_name)

            try:
                # Convert markdown to DOCX using pandoc
                with open('/tmp/convert_debug.log', 'a') as f:
                    f.write(f"DEBUG: About to run pandoc with temp_md_path={temp_md_path}, output_path={output_path}\n")
                    f.write(f"DEBUG: Temp file exists: {os.path.exists(temp_md_path)}\n")
                    f.write(f"DEBUG: Temp file size: {os.path.getsize(temp_md_path) if os.path.exists(temp_md_path) else 'N/A'}\n")

                result = subprocess.run(
                    ['pandoc', temp_md_path, '-f', 'markdown', '-t', 'docx', '-o', str(output_path)],
                    capture_output=True,
                    text=True,
                    check=True
                )

                with open('/tmp/convert_debug.log', 'a') as f:
                    f.write(f"DEBUG: Pandoc completed. Return code: {result.returncode}\n")
                    f.write(f"DEBUG: Pandoc stdout: {result.stdout}\n")
                    f.write(f"DEBUG: Pandoc stderr: {result.stderr}\n")
                    f.write(f"DEBUG: Output file exists: {os.path.exists(output_path)}\n")
                    f.write(f"DEBUG: Output file size: {os.path.getsize(output_path) if os.path.exists(output_path) else 'N/A'}\n")

                # Clean up temporary file
                os.unlink(temp_md_path)

                # Update resume_detail with the new filename
                update_query = text("""
                    UPDATE resume_detail
                    SET rewrite_file_name = :file_name
                    WHERE resume_id = :resume_id
                """)

                db.execute(update_query, {
                    "file_name": new_file_name,
                    "resume_id": resume_id
                })
                db.commit()

                return {"file_name": new_file_name}

            except subprocess.CalledProcessError as e:
                # Clean up temporary file on error
                if os.path.exists(temp_md_path):
                    os.unlink(temp_md_path)
                raise Exception(f"Pandoc conversion failed: {e.stderr}")

        except ValueError:
            raise
        except Exception as e:
            raise Exception(f"Error converting Markdown to DOCX: {str(e)}")

    # Markdown to ODT conversion
    @classmethod
    def md2odt(cls, markdown_content: str, output_file_name: str, reference_file: Optional[str] = None) -> str:
        """
        Convert Markdown content to ODT format.

        Args:
            markdown_content: Markdown text content
            output_file_name: Name for the output ODT file
            reference_file: Optional reference ODT file for styling (filename only, not full path)

        Returns:
            Path to the created ODT file
        """
        try:
            import subprocess
            import tempfile

            # Create temporary markdown file
            with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as temp_md:
                temp_md.write(markdown_content)
                temp_md_path = temp_md.name

            # Output path for odt
            output_path = cls._get_file_path(output_file_name)
            # Ensure directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            try:
                # Build pandoc command
                pandoc_cmd = ['pandoc', temp_md_path, '-f', 'markdown', '-t', 'odt', '-o', str(output_path)]

                # Add reference doc if provided
                if reference_file:
                    reference_path = cls._get_file_path(reference_file)
                    if os.path.exists(reference_path):
                        pandoc_cmd.extend(['--reference-doc', str(reference_path)])

                # Convert markdown to ODT using pandoc
                result = subprocess.run(
                    pandoc_cmd,
                    capture_output=True,
                    text=True,
                    check=True
                )

                # Clean up temporary file
                os.unlink(temp_md_path)

                return str(output_path)

            except subprocess.CalledProcessError as e:
                # Clean up temporary file on error
                if os.path.exists(temp_md_path):
                    os.unlink(temp_md_path)
                raise Exception(f"Pandoc conversion to ODT failed: {e.stderr}")

        except Exception as e:
            raise Exception(f"Error converting Markdown to ODT: {str(e)}")

    # Markdown to PDF conversion
    @classmethod
    def md2pdf(cls, markdown_content: str, output_file_name: str, output_dir: str = None) -> str:
        """
        Convert Markdown content to PDF format.

        Args:
            markdown_content: Markdown text content
            output_file_name: Name for the output PDF file
            output_dir: Optional directory path for output (defaults to RESUME_DIR)

        Returns:
            Path to the created PDF file
        """
        try:
            import subprocess
            import tempfile

            # Create temporary markdown file
            with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as temp_md:
                temp_md.write(markdown_content)
                temp_md_path = temp_md.name

            # Determine output path
            if output_dir:
                output_path = Path(output_dir) / output_file_name
                # Create directory if it doesn't exist
                os.makedirs(output_dir, exist_ok=True)
            else:
                output_path = cls._get_file_path(output_file_name)
                # Ensure directory exists
                os.makedirs(os.path.dirname(output_path), exist_ok=True)

            try:
                # Convert markdown to PDF using pandoc
                result = subprocess.run(
                    ['pandoc', temp_md_path, '-f', 'markdown', '-t', 'pdf', '-o', str(output_path)],
                    capture_output=True,
                    text=True,
                    check=True
                )

                # Clean up temporary file
                os.unlink(temp_md_path)

                return str(output_path)

            except subprocess.CalledProcessError as e:
                # Clean up temporary file on error
                if os.path.exists(temp_md_path):
                    os.unlink(temp_md_path)
                raise Exception(f"Pandoc conversion to PDF failed: {e.stderr}")

        except Exception as e:
            raise Exception(f"Error converting Markdown to PDF: {str(e)}")

    @classmethod
    def html2md(cls, html_content: str) -> str:
        """
        Convert resume HTML content to Markdown
        :param cls:
        :param html_content: contains the HTML content to be converted
        :return: a string containing the Markdown formatting of the document
        """

        from html_to_markdown import convert

        markdown = convert(html_content)
        return markdown


    # HTML to DOCX conversion (legacy method, keeping for backwards compatibility)
    @classmethod
    def html2docx_from_job(cls, job_id: int, db) -> dict:
        """
        Convert resume HTML to DOCX format for a specific job.

        This method:
        1. Retrieves the resume_html_rewrite for the job's associated resume
        2. Converts the HTML to DOCX format
        3. Saves the DOCX file with a -tailored.docx suffix
        4. Updates the resume_detail.rewrite_file_name field
        5. Returns the new filename

        Args:
            job_id: ID of the job
            db: Database session

        Returns:
            Dictionary with file_name key
        """
        try:
            from sqlalchemy import text

            # Query to get resume data
            query = text("""
                SELECT rd.resume_html_rewrite, rd.suggestion, r.file_name, r.resume_id, j.job_title, j.company 
                FROM job j
                JOIN resume r ON (j.resume_id = r.resume_id)
                JOIN resume_detail rd ON (r.resume_id = rd.resume_id)
                WHERE j.job_id = :job_id
            """)

            result = db.execute(query, {"job_id": job_id}).first()

            if not result:
                raise ValueError(f"No resume found for job_id {job_id}")

            if not result.resume_html_rewrite:
                raise ValueError(f"No rewritten HTML content found for job_id {job_id}")

            # Extract values
            html_content = result.resume_html_rewrite
            file_name = cls._set_file(result.company, result.job_title, 'docx')
            resume_id = result.resume_id

            logger.debug(f"Legacy HTML to DOCX conversion", job_id=job_id, html_length=len(html_content) if html_content else 0)

            # Check if HTML content is empty
            if not html_content or not html_content.strip():
                raise ValueError(f"No rewritten HTML content found for job_id {job_id}")

            # Convert HTML to DOCX using html4docx
            from html4docx import HtmlToDocx

            # Clean HTML
            #cleaned_html = cls._clean_html_for_docx(html_content)
            cleaned_html = html_content
            file_path = cls._get_file_path(file_name)

            try:
                logger.debug(f"Converting HTML to DOCX via html4docx")

                # Convert using html4docx
                from docx.shared import Inches

                parser = HtmlToDocx()
                doc = parser.parse_html_string(cleaned_html)

                # Set document margins to ensure content isn't clipped
                for section in doc.sections:
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.5)
                    section.left_margin = Inches(0.5)
                    section.right_margin = Inches(0.5)

                # Remove leading empty paragraphs
                while doc.paragraphs and not doc.paragraphs[0].text.strip():
                    p = doc.paragraphs[0]._element
                    p.getparent().remove(p)

                # Fix heading and paragraph formatting to prevent clipping
                from docx.shared import Pt
                from docx.enum.text import WD_LINE_SPACING

                '''
                #for paragraph in doc.paragraphs:
                    # Fix heading styles - ensure proper spacing and line height
                    #if paragraph.style.name.startswith('Heading'):
                        # Get the font size to calculate appropriate line spacing
                        #font_size = None
                        #if paragraph.runs:
                        #    for run in paragraph.runs:
                        #        if run.font.size:
                        #            font_size = run.font.size
                        #            break

                        # If no explicit font size, estimate based on heading level
                        #if font_size is None:
                        #    if 'Heading 1' in paragraph.style.name:
                        #        font_size = Pt(20)
                        #    elif 'Heading 2' in paragraph.style.name:
                        #        font_size = Pt(18)
                        #    elif 'Heading 3' in paragraph.style.name:
                        #        font_size = Pt(14)
                        #    else:
                        #        font_size = Pt(12)

                        # Reset paragraph spacing with generous room above
                        #paragraph.paragraph_format.space_before = Pt(24)
                        #paragraph.paragraph_format.space_after = Pt(12)

                        # Use MULTIPLE line spacing (1.5x) which scales with font size
                        #paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                        #paragraph.paragraph_format.line_spacing = 1.2

                        # Remove any indentation
                        #paragraph.paragraph_format.left_indent = Pt(0)
                        #paragraph.paragraph_format.first_line_indent = Pt(0)
                    #else:
                        # Fix regular paragraph indentation
                        #paragraph.paragraph_format.left_indent = Pt(0)
                        #paragraph.paragraph_format.first_line_indent = Pt(0)
                '''

                doc.save(str(file_path))

                logger.info(f"Converted using html4docx successfully")

            except Exception as e:
                logger.error(f"html4docx conversion failed: {str(e)}")
                raise

            # Update resume_detail with the cleaned HTML and new filename
            update_query = text("""
                UPDATE resume_detail
                SET resume_html_rewrite = :cleaned_html,
                    rewrite_file_name = :file_name
                WHERE resume_id = :resume_id
            """)

            db.execute(update_query, {
                "cleaned_html": cleaned_html,
                "file_name": file_name,
                "resume_id": resume_id
            })
            db.commit()

            logger.info(f"Updated resume_html_rewrite with cleaned HTML")

            return {"file_name": file_name}

        except ValueError:
            raise
        except Exception as e:
            raise Exception(f"Error converting HTML to DOCX: {str(e)}")


    @classmethod
    def _html_to_docx_direct(cls, html_content: str, output_path: Path) -> str:
        """
        Convert HTML to DOCX using python-docx directly for better control.

        Args:
            html_content: HTML string to convert
            output_path: Path where the DOCX file will be saved

        Returns:
            Cleaned HTML string (for database storage)
        """
        from bs4 import BeautifulSoup
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        # Clean the HTML first
        soup = BeautifulSoup(html_content, 'html.parser')

        # Remove problematic elements
        for tag in soup.find_all(['script', 'meta', 'title', 'style']):
            tag.decompose()

        # Remove wrapper divs with class attributes
        for div in soup.find_all('div', class_=True):
            div.unwrap()

        # Process diff tags
        for ins_tag in soup.find_all('ins'):
            ins_tag.unwrap()
        for del_tag in soup.find_all('del'):
            del_tag.decompose()

        # Create a new Document
        doc = Document()

        # Process the HTML elements
        def process_element(element, doc):
            """Recursively process HTML elements and add to document"""
            if element.name == 'h1':
                p = doc.add_heading(element.get_text().strip(), level=1)
            elif element.name == 'h2':
                p = doc.add_heading(element.get_text().strip(), level=2)
            elif element.name == 'h3':
                p = doc.add_heading(element.get_text().strip(), level=3)
            elif element.name == 'p':
                text = element.get_text().strip()
                if text:
                    # Check for text-based bullets in the paragraph
                    lines = text.split('\n')
                    bullet_lines = [line.strip() for line in lines if line.strip().startswith('- ')]
                    non_bullet_lines = [line.strip() for line in lines if line.strip() and not line.strip().startswith('- ')]

                    # Add non-bullet content first
                    if non_bullet_lines:
                        doc.add_paragraph('\n'.join(non_bullet_lines))

                    # Add bullets
                    for bullet in bullet_lines:
                        doc.add_paragraph(bullet[2:], style='List Bullet')
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text().strip(), style='List Bullet')
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text().strip(), style='List Number')
            elif element.name == 'hr':
                doc.add_paragraph('_' * 50)
            elif element.name in ['strong', 'b', 'em', 'i', 'a']:
                # These will be handled by parent paragraph
                pass
            elif hasattr(element, 'children'):
                for child in element.children:
                    if hasattr(child, 'name'):
                        process_element(child, doc)

        # Get body or root element
        body = soup.find('body')
        root = body if body else soup

        # Process all top-level elements
        for element in root.children:
            if hasattr(element, 'name'):
                process_element(element, doc)

        # Save the document
        doc.save(str(output_path))

        # Return cleaned HTML for database
        return str(soup)

    @classmethod
    def _clean_html_for_docx(cls, html_content: str) -> str:
        """
        Clean HTML using BeautifulSoup for proper parsing and structure.
        Preserves CSS styling and ensures proper HTML document structure.

        Args:
            html_content: Raw HTML content

        Returns:
            Cleaned, well-formed HTML string
        """
        from bs4 import BeautifulSoup

        # Parse the HTML - BeautifulSoup will fix malformed tags
        soup = BeautifulSoup(html_content, 'html.parser')

        # Remove problematic elements (keep style tags for formatting)
        for tag in soup.find_all(['script', 'meta', 'title']):
            tag.decompose()

        # Remove wrapper divs that have class attributes (from diff output)
        for div in soup.find_all('div', class_=True):
            # Unwrap the div (keep content, remove the div tag)
            div.unwrap()

        # Process diff tags (ins/del from htmldiff)
        # Replace <ins> with its text content (accepted additions)
        for ins_tag in soup.find_all('ins'):
            ins_tag.unwrap()

        # Remove <del> completely (rejected deletions)
        for del_tag in soup.find_all('del'):
            del_tag.decompose()

        # Remove empty list items (artifacts)
        for li in soup.find_all('li'):
            if not li.get_text(strip=True):
                li.decompose()

        # Check if HTML already has proper document structure
        has_html = soup.find('html') is not None
        has_body = soup.find('body') is not None

        if has_html and has_body:
            # Already has structure, just return the cleaned version
            return str(soup)
        elif has_body:
            # Has body but no html wrapper
            body_content = str(soup.body)
            return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
</head>
{body_content}
</html>'''
        else:
            # No document structure, wrap everything
            body_content = str(soup)
            return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
</head>
<body>
{body_content}
</body>
</html>'''

    # New HTML conversion methods for /convert/final endpoint
    @classmethod
    def html2docx(cls, html_content: str, output_filename: str) -> tuple[Path, str]:
        """
        Convert HTML content to DOCX format via WeasyPrint (HTML->PDF->DOCX).

        This approach uses WeasyPrint to generate a well-formatted PDF that honors CSS,
        then converts that PDF to DOCX format.

        Args:
            html_content: HTML string to convert
            output_filename: Name for the output DOCX file

        Returns:
            Tuple of (Path to created DOCX file, cleaned HTML string)

        Raises:
            Exception: If conversion fails
        """
        try:
            from html4docx import HtmlToDocx

            # Log the HTML content length for debugging
            logger.debug(f"HTML content length before cleaning", length=len(html_content) if html_content else 0)

            # Check if HTML content is empty or None
            if not html_content or not html_content.strip():
                raise Exception("HTML content is empty or None")

            # Clean HTML
            cleaned_html = cls._clean_html_for_docx(html_content)

            # Write DOCX file to resume directory
            file_path = cls._get_file_path(output_filename)

            # Ensure directory exists
            os.makedirs(os.path.dirname(file_path), exist_ok=True)

            try:
                logger.debug(f"Converting HTML to DOCX via html4docx")

                # Convert using html4docx
                from docx.shared import Inches

                parser = HtmlToDocx()
                doc = parser.parse_html_string(cleaned_html)

                # Set document margins to ensure content isn't clipped
                for section in doc.sections:
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.5)
                    section.left_margin = Inches(0.75)
                    section.right_margin = Inches(0.75)

                # Remove leading empty paragraphs
                while doc.paragraphs and not doc.paragraphs[0].text.strip():
                    p = doc.paragraphs[0]._element
                    p.getparent().remove(p)

                # Fix heading and paragraph formatting to prevent clipping
                from docx.shared import Pt
                from docx.enum.text import WD_LINE_SPACING

                for paragraph in doc.paragraphs:
                    # Fix heading styles - ensure proper spacing and line height
                    if paragraph.style.name.startswith('Heading'):
                        # Get the font size to calculate appropriate line spacing
                        font_size = None
                        if paragraph.runs:
                            for run in paragraph.runs:
                                if run.font.size:
                                    font_size = run.font.size
                                    break

                        # If no explicit font size, estimate based on heading level
                        if font_size is None:
                            if 'Heading 1' in paragraph.style.name:
                                font_size = Pt(16)
                            elif 'Heading 2' in paragraph.style.name:
                                font_size = Pt(14)
                            else:
                                font_size = Pt(12)

                        # Reset paragraph spacing with generous room above
                        paragraph.paragraph_format.space_before = Pt(24)
                        paragraph.paragraph_format.space_after = Pt(12)

                        # Use MULTIPLE line spacing (1.5x) which scales with font size
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                        paragraph.paragraph_format.line_spacing = 1.5

                        # Remove any indentation
                        paragraph.paragraph_format.left_indent = Pt(0)
                        paragraph.paragraph_format.first_line_indent = Pt(0)
                    else:
                        # Fix regular paragraph indentation
                        paragraph.paragraph_format.left_indent = Pt(0)
                        paragraph.paragraph_format.first_line_indent = Pt(0)

                doc.save(str(file_path))

                logger.info(f"HTML to DOCX conversion completed via html4docx", output_file=output_filename)

            except Exception as e:
                logger.error(f"html4docx conversion failed", error=str(e))
                raise

            return file_path, cleaned_html

        except Exception as e:
            logger.error(f"Error converting HTML to DOCX", error=str(e), output_filename=output_filename)
            raise Exception(f"Error converting HTML to DOCX: {str(e)}")


    @classmethod
    def html2pdf(cls, html_content: str, output_filename: str) -> Path:
        """
        Convert HTML content to PDF format using WeasyPrint.
        WeasyPrint honors CSS, so styles will be preserved.

        Args:
            html_content: HTML string to convert
            output_filename: Name for the output PDF file

        Returns:
            Path to the created PDF file

        Raises:
            Exception: If conversion fails
        """
        try:
            from weasyprint import HTML

            # Clean HTML but preserve CSS (WeasyPrint honors it)
            cleaned_html = cls._clean_html_for_docx(html_content)

            # Write PDF file to resume directory
            file_path = cls._get_file_path(output_filename)

            # Ensure directory exists
            os.makedirs(os.path.dirname(file_path), exist_ok=True)

            # Convert HTML to PDF (WeasyPrint honors CSS styles)
            HTML(string=cleaned_html).write_pdf(str(file_path))

            logger.info(f"HTML to PDF conversion completed", output_file=output_filename)
            return file_path

        except Exception as e:
            logger.error(f"Error converting HTML to PDF", error=str(e))
            raise Exception(f"Error converting HTML to PDF: {str(e)}")


    @classmethod
    def html2odt(cls, html_content: str, output_filename: str) -> Path:
        """
        Convert HTML content to ODT format using pandoc.

        Args:
            html_content: HTML string to convert
            output_filename: Name for the output ODT file

        Returns:
            Path to the created ODT file

        Raises:
            Exception: If conversion fails
        """
        try:
            import tempfile
            import subprocess

            # Create temporary HTML file
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as temp_html:
                temp_html.write(html_content)
                temp_html_path = temp_html.name

            # Get output file path
            file_path = cls._get_file_path(output_filename)

            # Ensure directory exists
            os.makedirs(os.path.dirname(file_path), exist_ok=True)

            try:
                # Use pandoc to convert HTML to ODT
                subprocess.run(
                    ['pandoc', temp_html_path, '-o', str(file_path)],
                    check=True,
                    capture_output=True,
                    text=True
                )

                logger.info(f"HTML to ODT conversion completed", output_file=output_filename)
                return file_path

            finally:
                # Clean up temporary HTML file
                if os.path.exists(temp_html_path):
                    os.unlink(temp_html_path)

        except subprocess.CalledProcessError as e:
            logger.error(f"Pandoc conversion failed", error=e.stderr)
            raise Exception(f"Error converting HTML to ODT: {e.stderr}")
        except Exception as e:
            logger.error(f"Error converting HTML to ODT", error=str(e))
            raise Exception(f"Error converting HTML to ODT: {str(e)}")

    # ========================================================================================
    # STANDARDIZED CONVERSION METHODS
    # All methods follow the signature: (input_path: str, output_path: str) -> bool
    # Returns True on success, False on failure
    # ========================================================================================
    @classmethod
    def docx2html_docx2html(cls, input_path: str, output_path: str) -> bool:
        """
        Convert Docx file to HTML using the docx2html library package

        Args:
            input_path: Full path to input DOCX file
            output_path: Full path for output HTML file

        Returns:
            True if conversion successful, False otherwise
        """

        from docx2html import convert

        html = convert(input_path)

        # Write to output file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)

        return True

    @classmethod
    def docx2html_docx_parser_converter(cls, input_path: str, output_path: str) -> bool:
        """
        Convert DOCX file to HTML using docx-parser-converter library.

        Args:
            input_path: Full path to input DOCX file
            output_path: Full path for output HTML file

        Returns:
            True if conversion successful, False otherwise
        """
        try:
            from docx_parser_converter.docx_to_html.docx_to_html_converter import DocxToHtmlConverter
            from docx_parser_converter.docx_parsers.utils import read_binary_from_file_path

            if not os.path.exists(input_path):
                logger.error(f"Input file not found", input_path=input_path)
                return False

            docx_file_content = read_binary_from_file_path(input_path)
            converter = DocxToHtmlConverter(docx_file_content, use_default_values=True)
            html_output = converter.convert_to_html()

            # Write to output file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_output)

            logger.info(f"DOCX to HTML conversion successful", input=input_path, output=output_path)
            return True

        except Exception as e:
            logger.error(f"DOCX to HTML conversion failed", error=str(e), input=input_path)
            return False

    @classmethod
    def odt2html_pandoc(cls, input_path: str, output_path: str) -> bool:
        """
        Convert ODT file to HTML using pandoc.

        Args:
            input_path: Full path to input ODT file
            output_path: Full path for output HTML file

        Returns:
            True if conversion successful, False otherwise
        """
        try:
            import subprocess

            if not os.path.exists(input_path):
                logger.error(f"Input file not found", input_path=input_path)
                return False

            # Use pandoc for ODT to HTML conversion
            result = subprocess.run(
                ['pandoc', input_path, '-f', 'odt', '-t', 'html', '-o', output_path],
                capture_output=True,
                text=True,
                check=True
            )

            logger.info(f"ODT to HTML conversion successful", input=input_path, output=output_path)
            return True

        except subprocess.CalledProcessError as e:
            logger.error(f"ODT to HTML conversion failed", error=e.stderr, input=input_path)
            return False
        except Exception as e:
            logger.error(f"ODT to HTML conversion failed", error=str(e), input=input_path)
            return False

    @classmethod
    def odt2docx_pandoc(cls, input_path: str, output_path: str) -> bool:
        """
        Convert ODT file to DOCX using pandoc.

        Args:
            input_path: Full path to input ODT file
            output_path: Full path for output DOCX file

        Returns:
            True if conversion successful, False otherwise
        """
        try:
            import subprocess

            if not os.path.exists(input_path):
                logger.error(f"Input file not found", input_path=input_path)
                return False

            # Use pandoc for ODT to DOCX conversion
            result = subprocess.run(
                ['pandoc', input_path, '-f', 'odt', '-t', 'docx', '-o', output_path],
                capture_output=True,
                text=True,
                check=True
            )

            logger.info(f"ODT to DOCX conversion successful", input=input_path, output=output_path)
            return True

        except subprocess.CalledProcessError as e:
            logger.error(f"ODT to DOCX conversion failed", error=e.stderr, input=input_path)
            return False
        except Exception as e:
            logger.error(f"ODT to DOCX conversion failed", error=str(e), input=input_path)
            return False

    @classmethod
    def pdf2html_markitdown(cls, input_path: str, output_path: str) -> bool:
        """
        Convert PDF file to HTML using MarkItDown (PDF -> Markdown -> HTML).

        Args:
            input_path: Full path to input PDF file
            output_path: Full path for output HTML file

        Returns:
            True if conversion successful, False otherwise
        """
        try:
            from markitdown import MarkItDown

            if not os.path.exists(input_path):
                logger.error(f"Input file not found", input_path=input_path)
                return False

            # Convert PDF to Markdown using MarkItDown
            md = MarkItDown()
            result = md.convert(input_path)
            markdown_content = result.text_content

            # Convert Markdown to HTML
            html_content = cls._markdown_to_html(markdown_content)

            # Write to output file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

            logger.info(f"PDF to HTML conversion successful", input=input_path, output=output_path)
            return True

        except Exception as e:
            logger.error(f"PDF to HTML conversion failed", error=str(e), input=input_path)
            return False

    @classmethod
    def html2docx_html4docx(cls, input_path: str, output_path: str) -> bool:
        """
        Convert HTML file to DOCX using html4docx library.

        Args:
            input_path: Full path to input HTML file
            output_path: Full path for output DOCX file

        Returns:
            True if conversion successful, False otherwise
        """
        try:
            from html4docx import HtmlToDocx
            from docx.shared import Inches

            if not os.path.exists(input_path):
                logger.error(f"Input file not found", input_path=input_path)
                return False

            # Read HTML content
            with open(input_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            # Clean HTML
            #cleaned_html = cls._clean_html_for_docx(html_content)
            cleaned_html = html_content

            # Convert using html4docx
            parser = HtmlToDocx()
            doc = parser.parse_html_string(cleaned_html)

            '''
            # Set document margins
            for section in doc.sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            # Fix heading clipping issues
            from docx.shared import Pt
            from docx.enum.text import WD_LINE_SPACING

            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    paragraph.paragraph_format.space_before = Pt(24)
                    paragraph.paragraph_format.space_after = Pt(12)
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    paragraph.paragraph_format.line_spacing = 1.2
            '''

            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            # Save document
            doc.save(output_path)

            logger.info(f"HTML to DOCX conversion successful", input=input_path, output=output_path)
            return True

        except Exception as e:
            logger.error(f"HTML to DOCX conversion failed", error=str(e), input=input_path)
            return False

    @classmethod
    def html2odt_pandoc(cls, input_path: str, output_path: str) -> bool:
        """
        Convert HTML file to ODT using pandoc.

        Args:
            input_path: Full path to input HTML file
            output_path: Full path for output ODT file

        Returns:
            True if conversion successful, False otherwise
        """
        try:
            import subprocess

            if not os.path.exists(input_path):
                logger.error(f"Input file not found", input_path=input_path)
                return False

            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            # Use pandoc to convert HTML to ODT
            subprocess.run(
                ['pandoc', input_path, '-o', output_path],
                check=True,
                capture_output=True,
                text=True
            )

            logger.info(f"HTML to ODT conversion successful", input=input_path, output=output_path)
            return True

        except subprocess.CalledProcessError as e:
            logger.error(f"HTML to ODT conversion failed", error=e.stderr, input=input_path)
            return False
        except Exception as e:
            logger.error(f"HTML to ODT conversion failed", error=str(e), input=input_path)
            return False

    @classmethod
    def html2pdf_weasyprint(cls, input_path: str, output_path: str) -> bool:
        """
        Convert HTML file to PDF using WeasyPrint.

        Args:
            input_path: Full path to input HTML file
            output_path: Full path for output PDF file

        Returns:
            True if conversion successful, False otherwise
        """
        try:
            from weasyprint import HTML

            if not os.path.exists(input_path):
                logger.error(f"Input file not found", input_path=input_path)
                return False

            # Read HTML content
            with open(input_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            # Clean HTML
            cleaned_html = cls._clean_html_for_docx(html_content)

            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            # Convert HTML to PDF
            HTML(string=cleaned_html).write_pdf(output_path)

            logger.info(f"HTML to PDF conversion successful", input=input_path, output=output_path)
            return True

        except Exception as e:
            logger.error(f"HTML to PDF conversion failed", error=str(e), input=input_path)
            return False

    # ========================================================================================
    # CONVERTAPI CONVERSION METHODS
    # Uses ConvertAPI service for conversions
    # ========================================================================================

    @classmethod
    def docx2html_convertapi(cls, input_path: str, output_path: str) -> bool:
        """
        Convert DOCX file to HTML using ConvertAPI service.

        Args:
            input_path: Full path to input DOCX file
            output_path: Full path for output HTML file

        Returns:
            True if conversion successful, False otherwise
        """
        try:
            import convertapi

            if not os.path.exists(input_path):
                logger.error(f"Input file not found", input_path=input_path)
                return False

            # Get ConvertAPI key from database
            api_key = cls._get_convertapi_key()
            if not api_key:
                return False

            # Ensure output directory exists before conversion
            output_dir = os.path.dirname(output_path)
            if output_dir:
                os.makedirs(output_dir, exist_ok=True)

            # Set API credentials (the library uses api_credentials, not api_secret)
            convertapi.api_credentials = api_key
            logger.debug(f"ConvertAPI conversion starting", input=input_path, output=output_path)

            # Convert file
            result = convertapi.convert('html', {'File': input_path}, from_format='docx')

            # Save result
            result.file.save(output_path)

            logger.info(f"DOCX to HTML conversion (ConvertAPI) successful", input=input_path, output=output_path)
            return True

        except Exception as e:
            logger.error(f"DOCX to HTML conversion (ConvertAPI) failed", error=str(e), input=input_path)
            import traceback
            logger.debug(f"ConvertAPI traceback", traceback=traceback.format_exc())
            return False

    @classmethod
    def html2docx_convertapi(cls, input_path: str, output_path: str) -> bool:
        """
        Convert HTML file to DOCX using ConvertAPI service.

        Args:
            input_path: Full path to input HTML file
            output_path: Full path for output DOCX file

        Returns:
            True if conversion successful, False otherwise
        """
        try:
            import convertapi

            if not os.path.exists(input_path):
                logger.error(f"Input file not found", input_path=input_path)
                return False

            # Get ConvertAPI key from database
            api_key = cls._get_convertapi_key()
            if not api_key:
                return False

            # Set API secret
            convertapi.api_credentials = api_key

            # Convert file
            result = convertapi.convert('docx', {'File': input_path}, from_format='html')

            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            # Save result
            result.file.save(output_path)

            logger.info(f"HTML to DOCX conversion (ConvertAPI) successful", input=input_path, output=output_path)
            return True

        except Exception as e:
            logger.error(f"HTML to DOCX conversion (ConvertAPI) failed", error=str(e), input=input_path)
            return False

    @classmethod
    def html2odt_convertapi(cls, input_path: str, output_path: str) -> bool:
        """
        Convert HTML file to ODT using ConvertAPI service.

        Args:
            input_path: Full path to input HTML file
            output_path: Full path for output ODT file

        Returns:
            True if conversion successful, False otherwise
        """
        try:
            import convertapi

            if not os.path.exists(input_path):
                logger.error(f"Input file not found", input_path=input_path)
                return False

            # Get ConvertAPI key from database
            api_key = cls._get_convertapi_key()
            if not api_key:
                return False

            # Set API secret
            convertapi.api_credentials = api_key

            # Convert file
            result = convertapi.convert('odt', {'File': input_path}, from_format='html')

            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            # Save result
            result.file.save(output_path)

            logger.info(f"HTML to ODT conversion (ConvertAPI) successful", input=input_path, output=output_path)
            return True

        except Exception as e:
            logger.error(f"HTML to ODT conversion (ConvertAPI) failed", error=str(e), input=input_path)
            return False

    @classmethod
    def html2pdf_convertapi(cls, input_path: str, output_path: str) -> bool:
        """
        Convert HTML file to PDF using ConvertAPI service.

        Args:
            input_path: Full path to input HTML file
            output_path: Full path for output PDF file

        Returns:
            True if conversion successful, False otherwise
        """
        try:
            import convertapi

            if not os.path.exists(input_path):
                logger.error(f"Input file not found", input_path=input_path)
                return False

            # Get ConvertAPI key from database
            api_key = cls._get_convertapi_key()
            if not api_key:
                return False

            # Set API secret
            convertapi.api_credentials = api_key

            # Convert file
            result = convertapi.convert('pdf', {'File': input_path}, from_format='html')

            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            # Save result
            result.file.save(output_path)

            logger.info(f"HTML to PDF conversion (ConvertAPI) successful", input=input_path, output=output_path)
            return True

        except Exception as e:
            logger.error(f"HTML to PDF conversion (ConvertAPI) failed", error=str(e), input=input_path)
            return False

    # ========================================================================================
    # ROUTING METHOD
    # Routes conversion requests to the appropriate method based on database settings
    # ========================================================================================

    @classmethod
    def convert_file(cls, source_format: str, target_format: str, input_path: str, output_path: str) -> bool:
        """
        Convert a file from source format to target format using the configured conversion method.

        This routing method reads conversion preferences from the personal table for HTML conversions
        and uses fixed methods for Markdown conversions.

        Args:
            source_format: Source file format (e.g., 'docx', 'odt', 'pdf', 'html', 'md')
            target_format: Target file format (e.g., 'html', 'docx', 'odt', 'pdf', 'md')
            input_path: Full path to input file
            output_path: Full path for output file

        Returns:
            True if conversion successful, False otherwise

        Supported conversions:
            - docx -> md (uses pandoc)
            - odt -> md (uses pandoc)
            - pdf -> md (uses markitdown)
            - docx -> html (methods: convertapi, docx-parser-converter)
            - odt -> html (methods: convertapi, pandoc)
            - pdf -> html (methods: convertapi, markitdown)
            - html -> docx (methods: convertapi, html4docx)
            - html -> odt (methods: convertapi, pandoc)
            - html -> pdf (methods: convertapi, weasyprint)
        """
        try:
            import subprocess
            from pathlib import Path

            # Normalize formats
            source_format = source_format.lower().strip()
            target_format = target_format.lower().strip()

            # Create conversion key (e.g., 'docx2html', 'html2docx')
            conversion_key = f"{source_format}2{target_format}"

            logger.debug(f"Convert file called", source_format=source_format, target_format=target_format, conversion_key=conversion_key)

            # Handle conversions to Markdown (not in personal table)
            if target_format == 'md':
                if source_format == 'docx':
                    # Use pandoc for DOCX to MD
                    result = subprocess.run(
                        ['pandoc', input_path, '-f', 'docx', '-t', 'markdown', '-o', output_path],
                        capture_output=True,
                        text=True,
                        check=True
                    )
                    return True
                elif source_format == 'odt':
                    # Use pandoc for ODT to MD
                    result = subprocess.run(
                        ['pandoc', input_path, '-f', 'odt', '-t', 'markdown', '-o', output_path],
                        capture_output=True,
                        text=True,
                        check=True
                    )
                    return True
                elif source_format == 'pdf':
                    # Use MarkItDown for PDF to MD
                    from markitdown import MarkItDown
                    md = MarkItDown()
                    result = md.convert(input_path)
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(result.text_content)
                    return True
                else:
                    logger.error(f"Unsupported conversion to markdown: {source_format}2md")
                    return False

            # Get conversion preference from database for HTML conversions
            from ..core.database import SessionLocal
            from sqlalchemy import text

            db = SessionLocal()
            try:
                query = text(f"SELECT {conversion_key} FROM personal LIMIT 1")
                result = db.execute(query).first()

                if not result:
                    logger.error("No personal settings found in database")
                    return False

                # Get the preferred conversion method
                preferred_method = getattr(result, conversion_key, None)
                logger.debug(f"preferred method", method=preferred_method)

            finally:
                db.close()

            # Default methods if no preference is set
            default_methods = {
                'docx2html': 'docx-parser-converter',
                'odt2html': 'pandoc',
                'pdf2html': 'markitdown',
                'html2docx': 'html4docx',
                'html2odt': 'pandoc',
                'html2pdf': 'weasyprint'
            }

            # Use default if no preference set
            if not preferred_method and conversion_key in default_methods:
                preferred_method = default_methods[conversion_key]

            if not preferred_method:
                logger.error(f"No conversion method configured for {conversion_key}")
                return False

            # Route to appropriate method based on conversion type and preferred method
            # If convertapi is selected and fails, automatically fall back to default method
            if conversion_key == 'docx2html':
                if preferred_method == 'convertapi':
                    result = cls.docx2html_convertapi(input_path, output_path)
                    if not result:
                        logger.warning(f"ConvertAPI failed for {conversion_key}, falling back to default method")
                        return cls.docx2html_docx_parser_converter(input_path, output_path)
                    return result
                elif preferred_method == 'docx-parser-converter':
                    return cls.docx2html_docx_parser_converter(input_path, output_path)

            elif conversion_key == 'odt2html':
                if preferred_method == 'convertapi':
                    result = cls.odt2html_convertapi(input_path, output_path)
                    if not result:
                        logger.warning(f"ConvertAPI failed for {conversion_key}, falling back to default method")
                        return cls.odt2html_pandoc(input_path, output_path)
                    return result
                elif preferred_method == 'pandoc':
                    return cls.odt2html_pandoc(input_path, output_path)

            elif conversion_key == 'pdf2html':
                if preferred_method == 'convertapi':
                    result = cls.pdf2html_convertapi(input_path, output_path)
                    if not result:
                        logger.warning(f"ConvertAPI failed for {conversion_key}, falling back to default method")
                        return cls.pdf2html_markitdown(input_path, output_path)
                    return result
                elif preferred_method == 'markitdown':
                    return cls.pdf2html_markitdown(input_path, output_path)

            elif conversion_key == 'html2docx':
                if preferred_method == 'convertapi':
                    result = cls.html2docx_convertapi(input_path, output_path)
                    if not result:
                        logger.warning(f"ConvertAPI failed for {conversion_key}, falling back to default method")
                        return cls.html2docx_html4docx(input_path, output_path)
                    return result
                elif preferred_method == 'html4docx':
                    return cls.html2docx_html4docx(input_path, output_path)

            elif conversion_key == 'html2odt':
                if preferred_method == 'convertapi':
                    result = cls.html2odt_convertapi(input_path, output_path)
                    if not result:
                        logger.warning(f"ConvertAPI failed for {conversion_key}, falling back to default method")
                        return cls.html2odt_pandoc(input_path, output_path)
                    return result
                elif preferred_method == 'pandoc':
                    return cls.html2odt_pandoc(input_path, output_path)

            elif conversion_key == 'html2pdf':
                if preferred_method == 'convertapi':
                    result = cls.html2pdf_convertapi(input_path, output_path)
                    if not result:
                        logger.warning(f"ConvertAPI failed for {conversion_key}, falling back to default method")
                        return cls.html2pdf_weasyprint(input_path, output_path)
                    return result
                elif preferred_method == 'weasyprint':
                    return cls.html2pdf_weasyprint(input_path, output_path)

            else:
                logger.error(f"Unsupported conversion: {conversion_key}")
                return False

            logger.error(f"Unknown conversion method: {preferred_method} for {conversion_key}")
            return False

        except subprocess.CalledProcessError as e:
            logger.error(f"File conversion failed",
                        source=source_format,
                        target=target_format,
                        error=e.stderr if hasattr(e, 'stderr') else str(e))
            return False
        except Exception as e:
            logger.error(f"File conversion routing failed",
                        source=source_format,
                        target=target_format,
                        error=str(e))
            return False
